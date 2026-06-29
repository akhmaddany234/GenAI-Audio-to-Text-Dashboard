# src/exporter.py

from docx import Document
from datetime import datetime


class DocumentExporter:

    def export(
        self,
        content,
        filename,
        title
    ):

        doc = Document()

        # Judul
        doc.add_heading(
            title,
            level=1
        )

        # Tanggal
        doc.add_paragraph(
            f"Tanggal: {datetime.now().strftime('%d-%m-%Y %H:%M')}"
        )

        doc.add_paragraph(
            "Dibuat secara otomatis menggunakan Generative AI."
        )

        doc.add_paragraph()

        # Isi dokumen
        for line in content.split("\n"):

            line = line.strip()

            if not line:
                continue

            if line.startswith("#"):
                heading = line.replace("#", "").strip()

                doc.add_heading(
                    heading,
                    level=2
                )

            else:
                doc.add_paragraph(line)

        doc.save(filename)

        return filename